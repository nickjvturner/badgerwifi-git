#!/usr/bin/env python3

"""
Written by Nick Turner (@nickjvturner@mastodon.social)
This script will find 'the' word doc in the same directory as the script

search through all the table cells and replace text strings from text_replacement.py
after this the table from the docx file will be converted into an xlsx file
"""

from pathlib import Path
import time
import sys

from docx import Document

from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter

sys.path.append(str(Path.cwd().parent))

from badgerwifitools.common import text_search


def main():
	# Get local file with extension .docx
	for file in sorted(Path.cwd().iterdir()):
		# ignore files in directory containing output
		if (file.suffix == '.docx') and (all(char not in file.stem for char in ('$', 'OUTPUT'))):
			proceed = input(f'Proceed with file: {str(file.name)}? (YES/no)')
			if proceed == 'no':
				continue

			print('filename:', file.name)

			document = Document(file)

			for table in document.tables:
				for row in table.rows:
					for cell in row.cells:
						for paragraph in cell.paragraphs:
							for key in text_search.keys():
								if key in paragraph.text:
									paragraph.text = paragraph.text.replace(key, text_search[key])
									print(key, 'text replaced with', text_search[key])
									# print(key, text_search[key])

			# Create a new Excel workbook
			wb = Workbook()
			ws = wb.active

			# Set the sheet name (you can modify this line to set a custom name)
			ws.title = 'AP Tracker'

			# Initialize row and column counters
			row_num = 1
			col_num = 1

			# Iterate through tables in the Word document
			for table in document.tables:
				for row in table.rows:
					for cell in row.cells:
						# Add cell content to the Excel worksheet
						ws.cell(row=row_num, column=col_num, value=cell.text)
						col_num += 1
					col_num = 1  # Reset column counter for each row
					row_num += 1

			# Set font size and make row 1 bold
			bold_font = Font(bold=True, size=16)
			for row in ws.iter_rows(min_row=1, max_row=1):
				for cell in row:
					cell.font = bold_font

			# Auto-fit all columns
			for column in ws.columns:
				max_length = 0
				column_letter = get_column_letter(column[0].column)
				for cell in column:
					try:
						if len(str(cell.value)) > max_length:
							max_length = len(cell.value)
					except Exception:
						pass
				adjusted_width = (max_length + 8)
				ws.column_dimensions[column_letter].width = adjusted_width

			# Save the Excel workbook
			wb.save(file.stem + ' - OUTPUT.xlsx')


if __name__ == "__main__":
	start_time = time.time()
	main()
	run_time = time.time() - start_time
	print(f'** Time to run: {round(run_time, 2)} seconds **')
