#!/usr/bin/env python3

"""
Written by Nick Turner (@nickjvturner@mastodon.social)
This script will find 'the' word doc in the same directory as the script
iterate through all the table cells and replace text strings with
corresponding images at specified heights.

Initially created in December 2022 in response to a throw-away comment by Jerry Olla @jolla
Latest meaningful update: 2023-02-14
"""

from pathlib import Path
import time

import inflect
from docx import Document
from docx.shared import Mm

nl = '\n'


def pretty_time_delta(seconds, lang=inflect.engine()):
	if not seconds:
		return f'0 seconds'
	seconds = int(seconds)
	days, seconds = divmod(seconds, 86400)
	hours, seconds = divmod(seconds, 3600)
	minutes, seconds = divmod(seconds, 60)
	measures = (
		(days, 'day'),
		(hours, 'hour'),
		(minutes, 'minute'),
		(seconds, 'second'),
	)
	return lang.join(
		[f'{count} {lang.plural(noun, count)}' for (count, noun) in measures if count]
	)


def main():
	# Get local file with extension .docx
	for file in sorted(Path.cwd().iterdir()):
		# ignore files in directory containing output
		if (file.suffix == '.docx') and (not ('OUTPUT' in file.stem)):
			proceed = input(f'{nl}Proceed with file: {str(file.name)}? (YES/no)')
			if proceed == 'no':
				exit()

			print(f'{nl}filename: {file.name}')

			document = Document(file)

			# print(document.tables)
			# for tables in document.tables:
			# 	print(tables)
			# 	for rows in tables.rows:
			# 		print(rows)
			# 		for cells in rows.cells:
			# 			print(cells.text)

			# Search terms are prefixed with string 'REPLACE-WITH-IMAGE='
			str_prefix = 'REPLACE-WITH-IMAGE='

			image_search = {
				'B': {
					'image': 'MNT-B.png',
					'height': 23},
				'H2': {
					'image': 'MNT-H2.png',
					'height': 23},
				'H3': {
					'image': 'MNT-H3.png',
					'height': 23},
				'V1A': {
					'image': 'MNT-V1A.png',
					'height': 23},

				'0.0°': {
					'image': 'EKA-TILT-0.png',
					'height': 18},
				'-20.0°': {
					'image': 'EKA-TILT-20.png',
					'height': 18}
			}

			image_insertion_points = []

			# count image insertions to be conducted
			for table in document.tables:
				for row in table.rows:
					for cell in row.cells:
						for paragraph in cell.paragraphs:
							for key in image_search.keys():
								if str_prefix + key in paragraph.text:
									image_insertion_points.append(key)
									print(len(image_insertion_points), end='\r')

			total_image_insertion_points = len(image_insertion_points)
			print(f'{nl}* {total_image_insertion_points} image insertion points identified *{nl}')

			# progress counter
			images_inserted = 1

			for table in document.tables:
				for row in table.rows:
					for cell in row.cells:
						for paragraph in cell.paragraphs:
							for key in image_search.keys():
								if str_prefix + key in paragraph.text:
									paragraph.text = paragraph.text.replace(str_prefix + key, "")
									picture_path = Path.cwd() / 'insert-images' / image_search[key]['image']
									paragraph.add_run().add_picture(str(picture_path), height=Mm(image_search[key]['height']))
									message = f"{key} image inserted, with height: {image_search[key]['height']} mm  ({images_inserted}/{total_image_insertion_points})"
									print(f'{message :>50}')
									images_inserted += 1

			# shutil.copy(file, Path.cwd() / Path(file.stem + '-ORIGINAL.docx'))
			document.save(Path.cwd() / Path(file.stem + '-OUTPUT-IMAGES_ADDED.docx'))

			print(f'{nl}* {images_inserted} images inserted *')


if __name__ == "__main__":
	start_time = time.time()
	main()
	run_time = time.time() - start_time
	print(f'** Time to run: {pretty_time_delta(run_time)} **')
