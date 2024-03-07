#!/usr/bin/env python3

"""
Written by Nick Turner (@nickjvturner@mastodon.social)

This script will find 'the' word doc in the same directory as the script
iterate through all the table cells and replace text strings with
corresponding images at specified heights.

Initially created in December 2022 in response to a throw-away comment by Jerry Olla @jolla
Latest meaningful update: 2023-10-07
"""

from pathlib import Path
import inflect
import time
import sys

from docx import Document
from docx.shared import Mm

# Get the current working directory as a Path object
cwd = Path.cwd()

# Iterate through parent directories
for parent in cwd.parents:
	# Construct the path to the common_dir.py file in the current parent directory
	common_path = parent / "text_replacement.py"

	# Check if the common_dir.py file exists
	if common_path.is_file():
		common_dir = str(parent) # Convert the path to a string and add the directory to sys.path
		sys.path.append(common_dir)

from badgerwifitools.common import image_search, text_search

nl = '\n'


def pretty_time_delta(seconds, lang=inflect.engine()):
	if not seconds:
		return f'0 seconds'
	seconds = int(seconds)
	days, seconds = divmod(seconds, 86400)
	hours, seconds = divmod(seconds, 3600)
	minutes, seconds = divmod(seconds, 60)
	measures = (
		(days, 'day'),
		(hours, 'hour'),
		(minutes, 'minute'),
		(seconds, 'second'),
	)
	return lang.join(
		[f'{count} {lang.plural(noun, count)}' for (count, noun) in measures if count]
	)


def main():
	docx_files = []  # Create empty list to store discovered docx files

	for file in sorted(Path.cwd().iterdir()):  # Get local file with extension .docx
		# ignore docx files in directory with $ and OUTPUT in filename
		if file.suffix == '.docx' and all(char not in file.stem for char in ('$', 'OUTPUT')):
			docx_files.append(file)  # Add file to the docx_files list

	for file in docx_files:
		print(file.name)

	proceed = input(f'{nl}Proceed with {len(docx_files)} files? (YES/no)')
	if proceed == 'no':
		exit()

	for file in docx_files:
		print(f'{nl}filename: {file.name}')

		document = Document(file)

		# scan AP placement maps and zoomed AP images
		annotated_AP_map_dir = Path.cwd() / 'insert-images' / 'zoomed_AP_workflow' / 'annotated'
		zoomed_AP_dir = Path.cwd() / 'insert-images' / 'zoomed_AP_workflow' / 'zoomed-APs'

		annotated_AP_map_dict = {}  # Create empty dictionary to store annotated AP map references
		annotated_AP_map_dict['height'] = 170  # add value for image insertion height
		for annotated_AP_map in annotated_AP_map_dir.iterdir():
			# print(annotated_AP_map)
			annotated_AP_map_dict[annotated_AP_map.stem] = annotated_AP_map


		zoomed_AP_dict = {}  # Create empty dictionary to store zoomed AP image references
		zoomed_AP_dict['height'] = 165
		for zoomed_AP in zoomed_AP_dir.iterdir():
			# print(zoomed_AP)
			zoomed_AP_dict[zoomed_AP.stem] = zoomed_AP


		# Search terms are prefixed with string 'REPLACE-WITH-IMAGE='
		str_prefix = 'REPLACE-WITH-IMAGE='

		# Keep track
		image_insertion_points = 0

		# count image insertions to be conducted
		for table in document.tables:
			for row in table.rows:
				count_image_insertion_point = True
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						if paragraph.text in ['Tilt:', 'Height:']:
							# if the row contains 'Tilt' or 'Height' ignore, skip it
							count_image_insertion_point = False
						for key in image_search.keys():
							if str_prefix + key in paragraph.text:
								if count_image_insertion_point:
									image_insertion_points += 1
									print('*', image_insertion_points, end='\r')

		total_image_insertion_points = int(image_insertion_points)

		print(f'{nl}* {total_image_insertion_points} image insertion points identified *{nl}')

		# progress counter
		images_inserted = 0

		# insert reference images
		for table in document.tables:
			for row in table.rows:
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						for key in image_search.keys():
							if str_prefix + key in paragraph.text:
								paragraph.text = paragraph.text.replace(str_prefix + key, "")
								picture_path = Path.cwd() / 'insert-images' / image_search[key]['image']
								paragraph.add_run().add_picture(str(picture_path), height=Mm(image_search[key]['height']))
								images_inserted += 1
								message = f"{key} image inserted, with height: {image_search[key]['height']} mm  ({images_inserted}/{total_image_insertion_points})"
								print(f'{message :>55}{nl}')

		# insert zoomed AP reference images
		for table in document.tables:
			for row in table.rows:
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						for key in zoomed_AP_dict.keys():
							if str_prefix + 'ZOOMED-AP-IMAGE=' + key in paragraph.text:
								paragraph.text = paragraph.text.replace(str_prefix + 'ZOOMED-AP-IMAGE=' + key, "")
								picture_path = zoomed_AP_dict[key]
								paragraph.add_run().add_picture(str(picture_path), height=Mm(zoomed_AP_dict['height']))
								images_inserted += 1
								# message = f"{key} image inserted, with height: {image_search[key]['height']} mm  ({images_inserted}/{total_image_insertion_points})"
								# print(f'{message :>55}{nl}')
								# print('yes')


		# insert AP placement maps
		for paragraph in document.paragraphs:
			# print(paragraph.text)
			for key in annotated_AP_map_dict.keys():
				print(str_prefix + 'CUSTOM-AP-LOCATION-MAP=' + key)
				if str_prefix + 'CUSTOM-AP-LOCATION-MAP=' + key in paragraph.text:
					# print(str_prefix + 'CUSTOM-AP-LOCATION-MAP=' + key)
					paragraph.text = paragraph.text.replace(str_prefix + 'CUSTOM-AP-LOCATION-MAP=' + key, "")
					picture_path = annotated_AP_map_dict[key]
					paragraph.add_run().add_picture(str(picture_path), height=Mm(annotated_AP_map_dict['height']))
					images_inserted += 1
					# message = f"{key} image inserted, with height: {image_search[key]['height']} mm  ({images_inserted}/{total_image_insertion_points})"
					# print(f'{message :>55}{nl}')
					print('yes')

		text_replacements = []

		for table in document.tables:
			for row in table.rows:
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						for key in text_search.keys():
							if key in paragraph.text:
								text_replacements.append(key)
								paragraph.text = paragraph.text.replace(key, text_search[key])
								print('*', len(text_replacements), 'text replacements', end='\r')
								time.sleep(0.01)

		print(f'{nl}------------{nl}Please wait while file is saved')
		document.save(Path.cwd() / Path(file.stem + '-OUTPUT-IMAGES_ADDED.docx'))

		print(f'{nl}* {images_inserted} images inserted *')
		print(f'* {len(text_replacements)} text strings replaced *')


if __name__ == "__main__":
	start_time = time.time()
	main()
	run_time = time.time() - start_time
	print(f'** Time to run: {pretty_time_delta(run_time)} **{nl}')
