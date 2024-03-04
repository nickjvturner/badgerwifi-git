#!/usr/bin/env python3

from pathlib import Path
import threading

import wx
from docx import Document
from docx.shared import Mm

from docx_manipulation.replacement_dict import image_search
from docx_manipulation.replacement_dict import text_search

nl = '\n'


def insert_images_threaded(docx_file, message_callback, progress_callback):
	# Wrapper function to run insert_images in a separate thread
	def run_in_thread():
		insert_images(docx_file, message_callback, progress_callback)
	# Update the UI or notify the user when done, if necessary
	# message_callback(f'Image Insertion process started')

	# Start the long-running task in a separate thread
	threading.Thread(target=run_in_thread).start()


def insert_images(docx_file, message_callback, progress_callback):
	wx.CallAfter(message_callback, f'Image Insertion process started')

	file = Path(docx_file)
	working_directory = file.parent
	print(f'{nl}filename: {file.name}')

	document = Document(str(file))

	# Search terms are prefixed with string 'REPLACE-WITH-IMAGE='
	str_prefix = 'REPLACE-WITH-IMAGE='

	# Keep track
	image_insertion_points = 0

	print(f'{nl}Searching file for image insertion points')
	wx.CallAfter(message_callback, f'Searching file for image insertion points')

	# count image insertions to be conducted
	for table in document.tables:
		for row in table.rows:
			count_image_insertion_point = True
			for cell in row.cells:
				for paragraph in cell.paragraphs:
					if paragraph.text in ['Tilt:', 'Height:']:
						# if the row contains 'Tilt' or 'Height' ignore, skip it
						count_image_insertion_point = False
					for key in image_search.keys():
						if str_prefix + key in paragraph.text:
							if count_image_insertion_point:
								image_insertion_points += 1
								wx.CallAfter(message_callback, f'counting image insertion strings: {image_insertion_points}')

	total_image_insertion_points = int(image_insertion_points)

	wx.CallAfter(message_callback, f'* {total_image_insertion_points} image insertion points identified *{nl}')

	# progress counter
	images_inserted = 0

	for table in document.tables:
		for row in table.rows:
			for cell in row.cells:
				for paragraph in cell.paragraphs:
					for key in image_search.keys():
						if str_prefix + key in paragraph.text:
							paragraph.text = paragraph.text.replace(str_prefix + key, "")
							picture_path = Path(__file__).resolve().parent / 'images' / image_search[key]['image']
							paragraph.add_run().add_picture(str(picture_path), height=Mm(image_search[key]['height']))
							images_inserted += 1
							message = f"{key} image inserted, with height: {image_search[key]['height']} mm  ({images_inserted}/{total_image_insertion_points})"
							wx.CallAfter(message_callback, message)
							wx.CallAfter(message_callback, f'images inserted: {images_inserted}/{total_image_insertion_points}')


	text_replacements = []

	wx.CallAfter(message_callback, f'{nl}Searching file for text replacements')

	for table in document.tables:
		for row in table.rows:
			for cell in row.cells:
				for paragraph in cell.paragraphs:
					for key in text_search.keys():
						if key in paragraph.text:
							text_replacements.append(key)
							paragraph.text = paragraph.text.replace(key, text_search[key])
							wx.CallAfter(message_callback, f'text replacements: {len(text_replacements)}')

	wx.CallAfter(message_callback, f'{nl}Please wait while file is saved{nl}')
	document.save(working_directory / Path(file.stem + '-OUTPUT-IMAGES_ADDED.docx'))

	message_callback(f'* {images_inserted} images inserted')
	message_callback(f'* {len(text_replacements)} text strings replaced *')
	message_callback(f'{nl}image insertion COMPLETE')
