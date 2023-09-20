#!/usr/bin/env python3

"""
Written by Nick Turner (@nickjvturner@mastodon.social)

This script will find 'the' word doc in the same directory as the script
iterate through all the table cells and replace text strings with
specified replacement strings.

Created December 2022 in response to a throw-away comment by Jerry Olla @jolla
"""

from pathlib import Path
from docx import Document
import time


def main():
	# Get local file with extension .docx
	for file in sorted(Path.cwd().iterdir()):
		# ignore files in directory containing output
		if (file.suffix == '.docx') and (not ('output' in file.stem)):
			proceed = input(f'Proceed with file: {str(file.name)}? (YES/no)')
			if proceed == 'no':
				exit()

			print('filename:', file.name)

			document = Document(file)

			text_search = {
				'Aruba AP-514 +  Aruba AP-ANT-45': 'Aruba AP-514',
				'Aruba AP-574 +  Aruba ANT-4x4-D608': 'Aruba AP-574',
				'Aruba AP-514 +  Aruba AP-ANT-48': 'Aruba AP-514',

				'Aruba AP-ANT-45 5GHz 5.5dBi': 'Aruba AP-ANT-45',
				'Aruba AP-ANT-48 5GHz 8.5dBi': 'Aruba AP-ANT-48',
				'Aruba ANT-3x3-D608 5GHz': 'Aruba ANT-3x3-D608',
				'Aruba ANT-4x4-D608 5GHz': 'Aruba ANT-4x4-D608',

				'Aruba AP-515 5GHz': 'Integrated',
				'Aruba AP-565 5GHz': 'Integrated',
				'Aruba AP-567 5GHz': 'Integrated',
				'Aruba AP-577 5GHz': 'Integrated',
				'Aruba AP-655 5GHz': 'Integrated',

				'0.0 dBm': ''
				}

			for table in document.tables:
				for row in table.rows:
					for cell in row.cells:
						for paragraph in cell.paragraphs:
							for key in text_search.keys():
								if key in paragraph.text:
									paragraph.text = paragraph.text.replace(key, text_search[key])
									print(key, 'text replaced with', text_search[key])
									# print(key, text_search[key])

			document.save(Path.cwd() / Path(file.stem + '-STRINGS SWAPPED.docx'))


if __name__ == "__main__":
	start_time = time.time()
	main()
	run_time = time.time() - start_time
	print(f'** Time to run: {round(run_time, 2)} seconds **')
