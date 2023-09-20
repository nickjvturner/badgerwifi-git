#!/usr/bin/env python3

"""
Written by Nick Turner (@nickjvturner@mastodon.social)
This script will find 'the' word doc in the same directory as the script
iterate through all the table cells and replace text strings with
specified replacement strings.

Created December 2022 in response to a throw-away comment by Jerry Olla @jolla
"""

from pathlib import Path
from docx import Document
import time

from common import text_search


def main():
	# Get local file with extension .docx
	for file in sorted(Path.cwd().iterdir()):
		# ignore files in directory containing output
		if (file.suffix == '.docx') and (all(char not in file.stem for char in ('$', 'SWAPPED'))):
			proceed = input(f'Proceed with file: {str(file.name)}? (YES/no)')
			if proceed == 'no':
				exit()

			print('filename:', file.name)

			document = Document(file)

			for table in document.tables:
				for row in table.rows:
					for cell in row.cells:
						for paragraph in cell.paragraphs:
							for key in text_search.keys():
								if key in paragraph.text:
									paragraph.text = paragraph.text.replace(key, text_search[key])
									print(key, 'text replaced with', text_search[key])
									# print(key, text_search[key])

			document.save(Path.cwd() / Path(file.stem + '-STRINGS SWAPPED.docx'))


if __name__ == "__main__":
	start_time = time.time()
	main()
	run_time = time.time() - start_time
	print(f'** Time to run: {round(run_time, 2)} seconds **')
